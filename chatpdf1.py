import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
import re
import time
import logging
from datetime import datetime
from typing import List, Dict, Tuple
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.cache import SQLiteCache
from langchain.globals import set_llm_cache
import google.generativeai as genai
from dotenv import load_dotenv
import pandas as pd
from io import BytesIO
from docx import Document as DocxDocument

# Configuration
load_dotenv()
set_llm_cache(SQLiteCache("cache.db"))
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
MAX_FILE_SIZE_MB = 25
SUPPORTED_FILE_TYPES = ["pdf", "docx", "txt"]
EMBEDDING_MODEL = "models/embedding-001"
DEFAULT_MODELS = ["gemini-1.5-flash-latest", "gemini-1.0-pro-latest"]

# Initialize Gemini
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

def validate_file(file) -> bool:
    """Validate file size and type."""
    if file.size > MAX_FILE_SIZE_MB * 1024 * 1024:
        st.error(f"File {file.name} exceeds {MAX_FILE_SIZE_MB}MB limit")
        return False
    extension = file.name.split(".")[-1].lower()
    if extension not in SUPPORTED_FILE_TYPES:
        st.error(f"Unsupported file type: {extension}")
        return False
    return True

def extract_text(file) -> Tuple[str, Dict]:
    """Extract text from various file formats with metadata tracking."""
    text = ""
    meta = {"source": file.name, "pages": []}
    
    try:
        if file.type == "application/pdf":
            pdf_reader = PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text() or ""
                text += page_text
                meta["pages"].append({
                    "page": page_num + 1,
                    "text": page_text[:500] + "..." if len(page_text) > 500 else page_text
                })
        
        elif file.type == "text/plain":
            text = str(file.read(), "utf-8")
            meta["pages"].append({"page": 1, "text": text[:500] + "..." if len(text) > 500 else text})
        
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = DocxDocument(BytesIO(file.read()))
            full_text = [para.text for para in doc.paragraphs]
            text = "\n".join(full_text)
            meta["pages"].append({"page": 1, "text": text[:500] + "..." if len(text) > 500 else text})
    
    except Exception as e:
        logger.error(f"Error processing {file.name}: {str(e)}")
        st.error(f"Error processing {file.name}: {str(e)}")
        return "", meta
    
    return text, meta

def process_files(files) -> Tuple[str, List[Dict]]:
    """Process multiple files with progress tracking."""
    all_text = []
    metadata = []
    
    # Use a main container for processing status
    status_container = st.empty()
    
    with status_container.status("Processing files...", expanded=True) as status:
        progress_bar = st.progress(0)
        total_files = len(files)
        
        for i, file in enumerate(files):
            if not validate_file(file):
                continue
                
            status.update(label=f"Reading {file.name}...")
            file_text, file_meta = extract_text(file)
            
            if file_text:
                all_text.append(file_text)
                metadata.append(file_meta)
            
            progress_bar.progress((i + 1) / total_files)
        
        status.update(label="Text extraction complete!", state="complete")
    
    return "\n\n".join(all_text), metadata

def chunk_text(text: str, metadata: List[Dict]) -> List[Document]:
    """Split text into chunks with metadata preservation."""
    if not text or not metadata:
        st.error("Error: No valid text extracted from documents")
        return []

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=200,
        length_function=len,
        add_start_index=True
    )
    
    chunks = text_splitter.split_text(text)
    if not chunks:
        st.error("Error: Text splitting produced no chunks")
        return []
    
    docs = []
    for i, chunk in enumerate(chunks):
        # Use the first document's metadata as a fallback
        source = metadata[0]["source"] if metadata else "unknown"
        doc = Document(
            page_content=chunk,
            metadata={
                "source": source,
                "page": 1,  # Default page number
                "chunk": i+1
            }
        )
        docs.append(doc)
    
    return docs

# Changed from @st.cache_data to @st.cache_resource to fix serialization error
@st.cache_resource(show_spinner=False)
def create_vector_store(_docs: List[Document]) -> FAISS:
    """Create vector store with caching."""
    if not _docs:
        st.error("No documents to process for vector store creation")
        return None

    try:
        embeddings = GoogleGenerativeAIEmbeddings(model=EMBEDDING_MODEL)
        return FAISS.from_documents(_docs, embeddings)
    except Exception as e:
        st.error(f"Error creating vector store: {str(e)}")
        logger.error(f"Vector store error: {str(e)}")
        return None

def get_qa_chain(model_name: str, temperature: float = 0.3):
    """Create QA chain with configurable parameters."""
    if not st.session_state.vector_store:
        st.error("Vector store not initialized")
        return None
        
    prompt_template = """
    Analyze the following context and question thoroughly. Provide a detailed, structured response with:
    1. Direct answer from context
    2. Supporting evidence
    3. Source document references
    4. Confidence estimation (0-100%)
    
    If unsure, state "Answer cannot be determined from provided documents."
    
    Context: {context}
    Question: {question}
    
    Response:
    """
    
    return RetrievalQA.from_chain_type(
        llm=ChatGoogleGenerativeAI(
            model=model_name,
            temperature=temperature,
            max_output_tokens=2000
        ),
        chain_type="stuff",
        retriever=st.session_state.vector_store.as_retriever(search_kwargs={"k": 5}),
        return_source_documents=True,
        chain_type_kwargs={"prompt": PromptTemplate.from_template(prompt_template)}
    )

def display_chat_history():
    """Render chat history with expandable details."""
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg["role"] == "assistant" and "sources" in msg:
                with st.expander("View Sources & Analysis"):
                    st.json(msg["sources"])

def handle_query(query: str):
    """Process user query with enhanced response handling."""
    if not query:
        return
        
    if "vector_store" not in st.session_state or not st.session_state.vector_store:
        st.error("Please process documents first")
        return
    
    start_time = time.time()
    
    with st.spinner("Analyzing documents..."):
        try:
            qa_chain = get_qa_chain(
                model_name=st.session_state.selected_model,
                temperature=st.session_state.temperature
            )
            
            if not qa_chain:
                return
                
            result = qa_chain.invoke({"query": query})
            
            response = {
                "content": result["result"],
                "sources": {
                    "documents": [
                        {
                            "source": doc.metadata.get("source", "unknown"),
                            "page": doc.metadata.get("page", 1),
                            "chunk": doc.metadata.get("chunk", 1),
                            "content": doc.page_content[:1000] + "..." if len(doc.page_content) > 1000 else doc.page_content
                        } for doc in result["source_documents"]
                    ],
                    "processing_time": f"{time.time() - start_time:.2f}s",
                    "model": st.session_state.selected_model,
                    "temperature": st.session_state.temperature
                }
            }
            
            st.session_state.messages.append({"role": "user", "content": query})
            st.session_state.messages.append({
                "role": "assistant",
                "content": response["content"],
                "sources": response["sources"]
            })
            
        except Exception as e:
            st.error(f"Analysis Error: {str(e)}")
            logger.error(f"Query Error: {str(e)}")

def model_selector():
    """Model selection and configuration sidebar."""
    st.sidebar.subheader("AI Configuration")
    model = st.sidebar.selectbox("Select Model", DEFAULT_MODELS, index=0)
    temp = st.sidebar.slider("Creativity Level", 0.0, 1.0, 0.3, 0.05)
    return model, temp

def file_analytics(metadata: List[Dict]):
    """Display file processing statistics."""
    st.sidebar.subheader("Document Analytics")
    
    if not metadata:
        st.sidebar.warning("No document metadata available")
        return
        
    try:
        df = pd.DataFrame([{
            "File": m["source"],
            "Pages": len(m.get("pages", [])),
            "Size": "Unknown"  # Cannot access file size after processing
        } for m in metadata])
        
        st.sidebar.dataframe(
            df,
            use_container_width=True,
            hide_index=True,
            column_config={"File": "Document", "Pages": "Pages", "Size": "Size"}
        )
    except Exception as e:
        st.sidebar.error(f"Error displaying analytics: {str(e)}")
    
    if "messages" in st.session_state and st.session_state.messages:
        st.sidebar.download_button(
            "Export Chat History",
            data=pd.DataFrame(st.session_state.messages).to_csv().encode("utf-8"),
            file_name=f"chat_history_{datetime.now().strftime('%Y%m%d%H%M')}.csv"
        )

def main():
    """Main application interface."""
    st.set_page_config(
        page_title="Enterprise Document AI",
        page_icon="📊",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    st.title("🔍 Intelligent Document Analysis System")
    st.markdown("### Multi-Format Document Processing with Advanced AI")
    
    # Initialize session state
    if "vector_store" not in st.session_state:
        st.session_state.vector_store = None
    if "metadata" not in st.session_state:
        st.session_state.metadata = []
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    # Configuration
    st.session_state.selected_model, st.session_state.temperature = model_selector()
    
    # File Processing - Simplified sidebar layout
    with st.sidebar:
        st.subheader("Document Upload")
        files = st.file_uploader(
            "Upload Documents",
            type=SUPPORTED_FILE_TYPES,
            accept_multiple_files=True,
            help=f"Max {MAX_FILE_SIZE_MB}MB per file"
        )
        
        if st.button("Process Documents"):
            if files:
                processed_text, metadata = process_files(files)
                if processed_text:
                    st.session_state.metadata = metadata
                    docs = chunk_text(processed_text, metadata)
                    if docs:
                        vector_store = create_vector_store(docs)
                        if vector_store:
                            st.session_state.vector_store = vector_store
                            st.success(f"Successfully processed {len(files)} documents")
                            st.rerun()
                        else:
                            st.error("Failed to create vector store")
                    else:
                        st.error("No valid document chunks created")
                else:
                    st.error("No text extracted from documents")
            else:
                st.warning("Please upload documents first")
        
        # Added Clear Cache button
        if st.button("Clear Cache"):
            st.cache_resource.clear()
            st.session_state.vector_store = None
            st.session_state.metadata = []
            st.session_state.messages = []
            st.success("Cache cleared successfully")
            st.rerun()
    
    # Main Interface
    col1, col2 = st.columns([3, 1])
    
    with col1:
        display_chat_history()
        
        if prompt := st.chat_input("Ask about your documents..."):
            handle_query(prompt)
            st.rerun()
    
    with col2:
        if st.session_state.metadata:
            file_analytics(st.session_state.metadata)
        
        st.subheader("Document Insights")
        if st.session_state.metadata:
            with st.expander("Document Previews"):
                for meta in st.session_state.metadata:
                    st.markdown(f"**{meta['source']}**")
                    page_count = len(meta.get('pages', []))
                    st.caption(f"Pages: {page_count}")
                    if page_count > 0 and 'text' in meta['pages'][0]:
                        st.text_area("Preview", meta['pages'][0]['text'], height=100)
        
        st.markdown("### Usage Tips")
        st.markdown("""
        - Ask complex questions using AND/OR operators
        - Request source verification for critical information
        - Use lower creativity for factual queries
        - Combine documents for cross-reference analysis
        """)

if __name__ == "__main__":
    main()